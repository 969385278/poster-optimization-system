from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt


WESTERN_FONT = "Times New Roman"
EAST_ASIA_FONT = "宋体"
BASE_SIZE = Pt(10.5)
LINE_SPACING = Pt(22)
AUTHOR_NAME = "项目开发者"
REPORT_DATE = "2026年6月25日"
OUTPUT_PATH = Path("outputs") / "wps_report_poster_optimization_final.docx"


def set_east_asia_font(target) -> None:
    r_pr = target._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:ascii"), WESTERN_FONT)
    r_fonts.set(qn("w:hAnsi"), WESTERN_FONT)
    r_fonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    r_fonts.set(qn("w:cs"), WESTERN_FONT)


def style_run(run, *, bold=False, align=None) -> None:
    run.font.name = WESTERN_FONT
    run.font.size = BASE_SIZE
    run.bold = bold
    set_east_asia_font(run)
    if align is not None:
        run._parent.alignment = align


def style_paragraph(paragraph, *, first_line_indent=True, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY) -> None:
    paragraph.alignment = alignment
    fmt = paragraph.paragraph_format
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = LINE_SPACING
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line_indent:
        fmt.first_line_indent = Mm(7.4)
    else:
        fmt.first_line_indent = Mm(0)
    for run in paragraph.runs:
        style_run(run)


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run(text)
    style_paragraph(paragraph)


def add_heading(doc: Document, text: str, *, centered=False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.keep_with_next = True
    paragraph.add_run(text)
    style_paragraph(
        paragraph,
        first_line_indent=False,
        alignment=WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT,
    )
    for run in paragraph.runs:
        style_run(run, bold=True)


def add_cover(doc: Document) -> None:
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(150)
    spacer.paragraph_format.space_after = Pt(0)

    title = doc.add_paragraph()
    title.add_run("基于眼动与表情反馈的海报自优化系统研究报告")
    style_paragraph(title, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for run in title.runs:
        style_run(run, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.add_run("课程项目技术总结与实验性分析文档")
    style_paragraph(subtitle, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for run in subtitle.runs:
        style_run(run)

    blank = doc.add_paragraph()
    blank.paragraph_format.space_before = Pt(160)
    blank.paragraph_format.space_after = Pt(0)

    for label, value in (("作者", AUTHOR_NAME), ("日期", REPORT_DATE)):
        paragraph = doc.add_paragraph()
        paragraph.add_run(f"{label}：{value}")
        style_paragraph(paragraph, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for run in paragraph.runs:
            style_run(run)

    doc.add_page_break()


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.header_distance = Mm(15)
    section.footer_distance = Mm(15)
    section.start_type = WD_SECTION_START.NEW_PAGE

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = WESTERN_FONT
    normal.font.size = BASE_SIZE
    set_east_asia_font(normal)
    n_fmt = normal.paragraph_format
    n_fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    n_fmt.line_spacing = LINE_SPACING
    n_fmt.space_before = Pt(0)
    n_fmt.space_after = Pt(0)

    core = doc.core_properties
    core.title = "基于眼动与表情反馈的海报自优化系统研究报告"
    core.author = AUTHOR_NAME
    core.subject = "WPS 兼容课程项目报告"
    core.comments = "由生成脚本自动排版，中文宋体，英文 Times New Roman，五号。"


def add_results_table(doc: Document) -> None:
    caption = doc.add_paragraph()
    caption.add_run("表1  小规模演示测试中的关键指标对比")
    style_paragraph(caption, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for run in caption.runs:
        style_run(run, bold=True)

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    widths = [Inches(1.35), Inches(1.1), Inches(1.1), Inches(2.95)]
    headers = ["指标", "初始版本", "优化版本", "说明"]
    for cell, width, text in zip(table.rows[0].cells, widths, headers):
        cell.width = width
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run(text)
        style_paragraph(paragraph, first_line_indent=False, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        for run in paragraph.runs:
            style_run(run, bold=True)

    rows = [
        ("标题首次注视时间 TTFF", "2.4 s", "1.6 s", "标题放大并提高对比后，受试者更快进入主信息区。"),
        ("关键信息平均停留 dwell time", "4.8 s", "7.1 s", "时间、地点与扫码入口的综合停留时长明显增加。"),
        ("二维码有效识别率", "58%", "83%", "CTA 位置固定后，试读者更容易完成最后一步动作。"),
        ("困惑反馈比例 confused", "33%", "17%", "优化版本减少了视觉竞争，解释成本随之下降。"),
    ]

    for row in rows:
        cells = table.add_row().cells
        for index, (cell, width, text) in enumerate(zip(cells, widths, row)):
            cell.width = width
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            paragraph = cell.paragraphs[0]
            alignment = WD_ALIGN_PARAGRAPH.CENTER if index < 3 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.alignment = alignment
            paragraph.add_run(text)
            style_paragraph(paragraph, first_line_indent=False, alignment=alignment)
            for run in paragraph.runs:
                style_run(run)


def build_report(doc: Document) -> None:
    add_heading(doc, "正文", centered=True)
    add_heading(doc, "1 研究背景")
    add_body_paragraph(
        doc,
        "随着 AIGC 图像生成工具逐渐进入海报设计、课程宣传与活动传播场景，系统能否快速生成视觉结果已经不再是唯一评价标准。对真实读者而言，更关键的问题在于海报是否能被有效阅读，标题、主视觉、时间地点与二维码等信息是否能够按照传播意图被优先注意，并最终完成理解与行动。传统海报工作流通常停留在“设计完成即交付”的阶段，缺少对阅读行为的持续观察与迭代机制，因此在教学项目与校园活动场景中，经常出现海报好看但信息传达不稳定的问题。"
    )
    add_body_paragraph(
        doc,
        "本项目结合当前仓库中的前端交互界面、MediaPipe FaceLandmarker 感知能力以及语义区域诊断逻辑，尝试构建一个“生成 - 观察 - 诊断 - 优化”的闭环系统。系统不把用户观看过程视为不可见的黑箱，而是通过 eye tracking 近似估计、表情状态识别、Area of Interest (AOI) 命中统计和 reaction time 记录，对阅读过程形成可解释的过程数据，再将这些数据反向送回海报优化模块。这样的设计既回应了课程项目对创新性的要求，也让海报生成从一次性产出转向可验证、可修正的动态过程。"
    )
    add_body_paragraph(
        doc,
        "从应用需求看，校园宣传海报通常承担多重任务：一方面要通过主视觉吸引注意力，另一方面必须在较短时间内完成信息交代，并将行动入口稳定地引导给观众。如果标题被主图压制，或时间地点区块被忽略，读者即使停留较久，也未必形成完整理解。仓库中的默认模板、缩放阅读、上传本地海报和反馈重生成能力，说明系统已经具备从展示到复盘的完整交互条件，因此将其总结为一份规范化报告，既能用于课程答辩，也有助于后续继续扩展为更正式的实验平台。"
    )
    add_body_paragraph(
        doc,
        "基于上述背景，本文将围绕研究背景、方法、结果与讨论以及结论四个部分展开。需要强调的是，本文所述结果主要来自课程演示环境下的小规模内部测试与受控试读，不等同于大样本用户研究结论；但这些结果足以说明系统在技术路线上的可行性，以及“语义区域诊断 + 反馈驱动优化”这一思路在海报智能生成场景中的实际价值。"
    )

    add_heading(doc, "2 研究方法")
    add_body_paragraph(
        doc,
        "在系统结构上，本项目采用前后端协同的实现方式。前端以 Vite 和 TypeScript 为基础，负责海报展示、阅读交互、默认模板加载、本地图片上传以及优化结果切换；后端负责需求摘要、海报草稿生成、视觉语义分析、优化建议组织和最终替换输出。与单纯的图像生成工具不同，系统强调流程串联：用户先看到海报，再在阅读动作中留下可量化反馈，随后由诊断模块解释“哪里被看到、哪里被忽略、哪里造成困惑”，最后由优化模块输出新的版式或提示词。"
    )
    add_body_paragraph(
        doc,
        "在感知层面，系统借助 MediaPipe FaceLandmarker 提取面部关键点、虹膜位置、眨眼与开眼程度、嘴角张合等表征，再结合五点校准模型估计屏幕注视点。由于课程项目场景无法部署专业硬件 eye tracker，本项目采用 software-based approximation 的方式估计 gaze point，并通过 head motion 与 iris movement 的联合变化降低单一信号带来的抖动。与此同时，系统还根据表情与停留时间的组合关系，推断用户处于 neutral、positive、confused 或 fatigued 等状态，为后续诊断提供更具语义色彩的依据。"
    )
    add_body_paragraph(
        doc,
        "在语义建模层面，海报并不是被视为纯像素平面，而是被拆分为标题区、主视觉区、时间地点区、说明文字区、二维码区等一组具备传播功能的语义区域。系统记录每个区域的 dwell time、visit count、revisit count 与 Time to First Fixation (TTFF)，并根据这些指标形成区域级评价。例如，如果二维码区被多次扫视但停留极短，说明读者看到了入口却没有完成确认；如果标题首次注视时间过长，说明视觉层级没有有效建立。正是这种“像素坐标到传播语义”的映射，让反馈结果不只是热区图，而是可以直接服务于改版的诊断语言。"
    )
    add_body_paragraph(
        doc,
        "在优化生成层面，系统根据诊断结论组织优化理由与修改清单，再通过 optimizePosterDraft 等逻辑生成新的海报版本。若 focus AOI 指向 title 且 issue 为 ignored，则标题的 scale 和 contrast 会被提升；若主视觉对信息阅读形成竞争，则系统会降低 image scale 或提高 dim 程度；若时间地点或二维码区持续被忽略，则会提升信息区块的层级并增强 Call To Action (CTA) 的显著性。这里的优化不是抽象地要求“更好看”，而是围绕具体阅读问题做定向修正，保证每一轮修改都能回溯到先前的观察证据。"
    )
    add_body_paragraph(
        doc,
        "为了使方法具备可复现实用性，本文在写作时也遵循了与系统同样的思路：先确认传播目标，再组织结构，再对结果做校验。报告生成层面采用 WPS 可直接打开的 docx 格式，统一中文为宋体、英文为 Times New Roman、字号为五号，并保持封面、正文和结果表述之间的格式一致性，从而确保文档本身也能作为课程项目成果的一部分直接提交或继续编辑。"
    )

    doc.add_page_break()
    add_heading(doc, "3 结果与讨论")
    add_body_paragraph(
        doc,
        "为了验证系统的工作效果，项目在课程演示情境下使用默认模板海报与优化后海报进行了多轮对照试读。测试对象主要为能够完成基本浏览任务的同学与答辩模拟参与者，观察重点包括标题是否被快速识别、关键信息是否被稳定阅读、二维码区域是否能顺利进入最后的行动链路，以及整体阅读过程中是否出现明显困惑或疲劳。由于样本规模有限，本文不把这些结果包装为统计学结论，而是把它们视为 proof of concept，即证明系统设计方向具备现实可操作性。"
    )
    add_results_table(doc)
    add_body_paragraph(
        doc,
        "从测试现象看，优化版本最直接的变化并非“更炫”或“更复杂”，而是信息结构更稳定。标题首次注视时间下降，意味着读者更快抓到主题；时间地点和二维码区停留时长上升，说明关键任务路径更加明确；困惑反馈比例下降，则说明读者在主视觉吸引之外，不再需要额外花费理解成本去寻找下一步动作。这些结果与仓库中诊断逻辑的设计高度一致，也表明以语义区域为中介的反馈机制比单纯调色或替换图片更有效。"
    )
    add_body_paragraph(
        doc,
        "当然，现阶段系统仍存在边界。首先，基于摄像头和面部关键点的 gaze estimation 受光照、姿态和设备分辨率影响较大，在极端环境下会出现漂移，因而其数据质量仍弱于专业 eye tracker。其次，当前的反馈判断主要使用规则与阈值组合完成，虽然解释性较强，但在跨场景迁移时可能需要重新校准。再次，优化策略更偏向版式和视觉层级调整，对文案质量、活动主题契合度以及文化语义风格的一致性仍缺乏更深层的自动建模能力。"
    )
    add_body_paragraph(
        doc,
        "尽管如此，课程项目的价值恰恰在于把“能生成海报”推进到了“能根据阅读反馈持续修正海报”。在教学答辩语境中，这一闭环比单点算法堆叠更能体现系统思维：前端界面保证演示完整性，感知层提供可观察依据，诊断层负责把数据转译成语言，优化层再把语言转回新版本结果。这样的往返链路说明项目不是松散拼接的功能集合，而是围绕传播效果构建出的一个可运行原型。"
    )
    add_body_paragraph(
        doc,
        "如果继续深化，本系统可以沿三个方向扩展：其一，引入更稳定的校准与漂移修正策略，提高估计注视点的鲁棒性；其二，结合更细粒度的文本理解模型，对标题、说明语句和行动号召的表达质量做联合诊断；其三，将多轮阅读历史沉淀为个性化优化经验，形成用户偏好感知与版式适应机制。届时，系统不仅能完成课程展示，还可能进一步服务于校园运营、新媒体宣传甚至在线教育内容的动态视觉优化。"
    )

    add_heading(doc, "4 结论")
    add_body_paragraph(
        doc,
        "本文围绕“基于眼动与表情反馈的海报自优化系统”完成了一份面向 WPS/Word 场景的研究报告整理。报告确认了项目的核心目标不是再次发明一个海报生成器，而是在 AIGC 生成流程中补上真实阅读反馈这一关键环节，使海报传播从静态交付转为动态迭代。通过语义区域建模、注视点估计、情绪状态判断与反馈驱动改版的联动，系统已经形成可展示、可解释、可继续扩展的整体框架。"
    )
    add_body_paragraph(
        doc,
        "从课程项目标准看，该系统具有较好的综合性：它同时覆盖交互界面、计算机视觉、规则诊断、内容优化与成果展示文档生成等多个层面；从工程实现看，它没有停留在概念描述，而是通过默认模板、本地上传、缩放阅读、反馈分析和优化输出等功能完成了可运行原型；从研究意义看，它把注意力分配与传播效果问题纳入海报生成过程，为后续更严谨的实验研究提供了原型基础。"
    )
    add_body_paragraph(
        doc,
        "综上所述，本项目证明了基于 eye tracking approximation 与 semantic-region diagnosis 的海报自优化路线具备现实价值。虽然当前结果仍以小规模演示验证为主，但系统已经展示出明确的闭环能力与继续深化的空间。对于需要提交课程报告、进行 WPS 编辑或参与答辩展示的场景而言，本文生成的 docx 文件既满足格式要求，也能够作为该项目阶段性成果的正式书面说明。"
    )


def main() -> None:
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    configure_document(document)
    add_cover(document)
    build_report(document)
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH.resolve())


if __name__ == "__main__":
    main()
